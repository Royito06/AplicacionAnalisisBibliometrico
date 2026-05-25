# Cálculos (Top 10, tasas de crecimiento)

import pandas as pd
import datetime
import io 
from docx import Document
import networkx as nx     

#--------------------------------------------------------------------------------------------

def obtener_rango_anios(df):
    """
    Busca la columna de años en el DataFrame y devuelve el año más antiguo y el más reciente.
    """
    posibles_nombres = ['Año', 'Year', 'PY', 'Publication Year', 'año', 'year']
    
    columna_anio = None

    for col in df.columns:
        if col in posibles_nombres:
            columna_anio = col
            break
            
    # Si el archivo no tiene columna de año, devolvemos un error controlado
    if columna_anio is None:
        return {"error": "No se encontró la columna de Año en el dataset."}
    
    # Eliminamos las filas que tengan esa celda vacía
    df_limpio = df.dropna(subset=[columna_anio])
    
    anio_minimo = int(df_limpio[columna_anio].min())
    anio_maximo = int(df_limpio[columna_anio].max())
    
    return {
        "minimo": anio_minimo,
        "maximo": anio_maximo,
        "mensaje_formateado": f"Periodo analizado: {anio_minimo} - {anio_maximo}"
    }

#--------------------------------------------------------------------------------------------

def calcular_promedio_publicaciones(df):
    """
    Calcula el promedio de publicaciones por autor dividiendo 
    el total de artículos entre el número de autores únicos.
    """
    posibles_nombres = ['Autor', 'Autores', 'Author', 'Authors', 'AU']
    col_autor = None
    
    for col in df.columns:
        if col in posibles_nombres:
            col_autor = col
            break
            
    if col_autor is None:
        return {"error": "No se encontró la columna de Autores en el dataset."}
    
    # Quitamos las filas que no tengan autor
    df_limpio = df.dropna(subset=[col_autor])
    total_publicaciones = len(df_limpio)
    
    # reemplazamos los ; por , y separamos cada nombre
    listas_de_autores = df_limpio[col_autor].astype(str).str.replace(';', ',').str.split(',')
    
    # Metemos a todos los autores y les quitamos los espacios extra
    todos_los_autores = [autor.strip() for sublista in listas_de_autores for autor in sublista if autor.strip()]
    
    #convertimos la lista en un set para eliminar a los repetidos y los contamos
    autores_unicos = len(set(todos_los_autores))

    if autores_unicos == 0:
        return {"error": "No hay autores válidos para calcular el promedio."}
        
    promedio = total_publicaciones / autores_unicos
    
    return {
        "total_publicaciones": total_publicaciones,
        "autores_unicos": autores_unicos,
        "promedio_calculado": round(promedio, 2),
        "mensaje_formateado": f"Productividad: {round(promedio, 2)} publicaciones por autor"
    }

#--------------------------------------------------------------------------------------------

def formatear_apa(fila, col_revista, col_volumen, col_numero, col_paginas):
    autores_sin_formato = str(fila.get('Authors', ''))
    autores_lista = [ a.strip() for a in autores_sin_formato.split(';') if a.strip()]
    if len(autores_lista) == 1:
        autores_apa = autores_lista[0]
    elif len(autores_lista)>1:
        autores_apa = ', '.join(autores_lista[:-1])+ ', & ' + autores_lista[-1]
    else:
        autores_apa = 'Autor desconocido'
        
    anio     = str(fila.get('Year',         ''))
    titulo   = str(fila.get('Title',        ''))
    revista  = str(fila.get(col_revista,    '')) if col_revista else ''
    volumen  = str(fila.get(col_volumen,    '')) if col_volumen else ''
    numero   = str(fila.get(col_numero,     '')) if col_numero  else ''
    paginas  = str(fila.get(col_paginas,    '')) if col_paginas else ''
    
    # Construir progresivamente — solo agrega lo que existe
    ref = f"{autores_apa}. ({anio}). {titulo}."
    if revista:
        ref += f" {revista}"
        if volumen:
            ref += f", {volumen}"
            if numero:
                ref += f"({numero})"
        if paginas:
            ref += f", {paginas}"
        ref += "."

    return ref


#--------------------------------------------------------------------------------------------

def obtener_top_10_autores(df):
    df = df.copy()
    posibles_nombres = ['Autor', 'Autores', 'Author', 'Authors', 'AU']
    col_autor = None

    for col in df.columns:
        if col in posibles_nombres:
            col_autor = col
            break

    if col_autor is None:
        return []

    df_limpio = df.dropna(subset=[col_autor])

    # Separar SOLO por ";" — la coma es parte del nombre (Apellido, Inicial)
    autores_individuales = (
        df_limpio[col_autor]
        .astype(str)
        .str.split(';')       # ← solo punto y coma, sin tocar las comas
        .explode()
        .str.strip()
    )
    
    
    autores_individuales = autores_individuales[autores_individuales != ""]

    top_10_serie = autores_individuales.value_counts().head(10)

    return [{"autor": autor, "cantidad": int(cantidad)} for autor, cantidad in top_10_serie.items()]

#--------------------------------------------------------------------------------------------

def obtener_top_10_trabajos(df):
    """
    Ordena los trabajos por número de citas y da el top 10
    """
    col_citas   = next((c for c in df.columns if 'cite' in c.lower()), None)
    col_titulo  = next((c for c in df.columns if 'titl' in c.lower()), None)
    col_revista = next((c for c in df.columns if 'source' in c.lower() or 'journal' in c.lower()), None)
    col_volumen = next((c for c in df.columns if c.lower() in ('volume', 'vol')), None)
    col_numero  = next((c for c in df.columns if c.lower() in ('issue', 'number')), None)
    col_paginas = next((c for c in df.columns if 'page' in c.lower()), None)

    if not col_citas or not col_titulo:
        return []

    df[col_citas] = pd.to_numeric(df[col_citas], errors='coerce').fillna(0)
    top_10 = df.sort_values(by=col_citas, ascending=False).head(10)

    resultados = []
    for _, fila in top_10.iterrows():
        resultados.append({
            "titulo": fila[col_titulo],
            "citas":  int(fila[col_citas]),
            "apa":    formatear_apa(fila, col_revista, col_volumen, col_numero, col_paginas)
        })
    return resultados
#--------------------------------------------------------------------------------------------

def contabilizar_coautorias(df):
    """
    Cuenta cuántas publicaciones tienen más de un autor
    """

    posibles_nombres = ['Autor', 'Autores', 'Author', 'Authors', 'AU']
    col_autor = None
    
    for col in df.columns:
        if col in posibles_nombres:
            col_autor = col
            break
            
    if col_autor is None:
        return {"error": "No se encontró la columna de Autores."}
        
    df_limpio = df.dropna(subset=[col_autor])
    total_publicaciones = len(df_limpio)
    
    # Separamos el texto de los autores en arreglos
    listas_de_autores = df_limpio[col_autor].astype(str).str.replace(';', ',').str.split(',')
    
    # Cuenta cuántas listas tienen más de 1 autor válido, cada lista y devuelve true si hay más de 1 autor/elemento
    es_coautorado = listas_de_autores.apply(lambda autores: len([a for a in autores if a.strip()]) > 1)
    
    total_coautoradas = int(es_coautorado.sum())
    porcentaje = 0
    if total_publicaciones > 0:
        porcentaje = (total_coautoradas / total_publicaciones) * 100
        
    return {
        "total_coautoradas": total_coautoradas,
        "porcentaje_coautoria": round(porcentaje, 2),
        "mensaje_formateado": f"{total_coautoradas} publicaciones en colaboración ({round(porcentaje, 2)}%)"
    }

#--------------------------------------------------------------------------------------------

def obtener_articulos_por_universidad(df, nombre_universidad):
    """
    Filtra el dataset para devolver solo los artículos que pertenezcan a una universidad específica.
    """
    # buscamos la columna de afiliación y la de título
    col_afil = None
    col_titulo = None
    
    col_afil   = next((c for c in df.columns if 'affil' in c.lower() or 'instit' in c.lower()), None)
    col_titulo = next((c for c in df.columns if 'titl' in c.lower()), None)
    
    

            
    if not col_afil or not col_titulo:
        return [] # vacío si faltan columnas
        
    # Quitamos filas sin institución
    df_limpio = df.dropna(subset=[col_afil])
    
    # Aplicamos el filtro para ignorar mayúsculas y minúsculas
    filtro = df_limpio[col_afil].str.contains(nombre_universidad, case=False, na=False)
    df_filtrado = df_limpio[filtro]
    
    resultados = []
    for _, fila in df_filtrado.iterrows():
        resultados.append({
            "titulo": fila[col_titulo],
            "afiliacion": fila[col_afil]
        })
        
    return resultados

#--------------------------------------------------------------------------------------------

def calcular_promedio_citas(df):
    """
    Calcula el promedio de citas por artículo sumando todas las citas 
    y dividiéndolas entre el total de publicaciones.
    """
    col_citas = None    
    col_citas = next((c for c in df.columns if 'cite' in c.lower() or 'cit' in c.lower()), None)

    

            
    if col_citas is None:
        return {"error": "No se encontró la columna de Citas en el dataset."}
        
    # Forzamos la conversión a números
    citas_numericas = pd.to_numeric(df[col_citas], errors='coerce').fillna(0)
    
    # Sumatoria
    total_citas = int(citas_numericas.sum())
    total_articulos = len(df)

    if total_articulos == 0:
        return {"error": "El dataset está vacío."}
        
    promedio = total_citas / total_articulos
    
    return {
        "total_citas": total_citas,
        "promedio_citas": round(promedio, 2),
        "mensaje_formateado": f"Impacto promedio: {round(promedio, 2)} citas por artículo"
    }

#--------------------------------------------------------------------------------------------

def obtener_lista_paises(df):
    # --- Paso 1: buscar columna de país explícita (WoS tiene 'Country') ---
    col_pais = next(
        (c for c in df.columns if 'countr' in c.lower() or 'país' in c.lower()),
        None
    )
    if col_pais:
        paises = (
            df[col_pais].dropna().astype(str)
            .str.split(';').explode().str.strip()
            .loc[lambda s: s != '']
            .value_counts()
        )
        return [{"pais": p, "cantidad": int(n)} for p, n in paises.items()]

    # --- Paso 2: extraer de Affiliations (Scopus) ---
    # Scopus usa "Affiliations" — buscar con nombre limpio
    col_afil = next(
        (c for c in df.columns if 'affil' in c.lower()),
        None
    )
    if col_afil is None:
        return []

    paises_extraidos = []

    for celda in df[col_afil].dropna().astype(str):
        # Cada celda tiene varias afiliaciones separadas por ";"
        for afiliacion in celda.split(';'):
            afiliacion = afiliacion.strip()
            if not afiliacion:
                continue
            # El país es el último fragmento separado por coma
            # pero hay que limpiar espacios, números y puntuación suelta
            partes = [p.strip() for p in afiliacion.split(',') if p.strip()]
            if len(partes) < 2:
                continue  # si solo hay un fragmento, no es útil
            candidato = partes[-1]
            # Descartar si es claramente un código postal o número
            if candidato.replace(' ', '').isdigit():
                candidato = partes[-2] if len(partes) >= 2 else None
            if candidato:
                paises_extraidos.append(candidato)

    if not paises_extraidos:
        return []

    serie = pd.Series(paises_extraidos).str.strip()
    # Normalizar variantes comunes
    normalizacion = {
        'United States': ['USA', 'U.S.A.', 'U.S.', 'United States of America'],
        'United Kingdom': ['UK', 'U.K.', 'England', 'Scotland', 'Wales'],
        'China': ["People's Republic of China", 'P.R. China', 'PR China'],
    }
    for nombre_correcto, variantes in normalizacion.items():
        serie = serie.replace(variantes, nombre_correcto)

    conteo = serie.value_counts()
    return [{"pais": p, "cantidad": int(n)} for p, n in conteo.items()]

#--------------------------------------------------------------------------------------------

def obtener_detalle_articulo(df, titulo_buscado):
    """
    Busca un artículo por su título exacto y devuelve todos sus metadatos en formato de diccionario.
    """
    posibles_titulos = ['Título', 'Title', 'TI', 'Article Title']
    col_titulo = None
    
    for col in df.columns:
        if col in posibles_titulos:
            col_titulo = col
            break
            
    if not col_titulo:
        return {"error": "No se encontró la columna de Títulos."}
        
    # Filtramos el dataframe buscando la fila donde el título coincida exactamente y
    # Usamos .astype(str) y .str.strip() por si hay espacios invisibles
    filtro = df[df[col_titulo].astype(str).str.strip() == titulo_buscado.strip()]
    
    if filtro.empty:
        return {"error": "No se encontró ningún artículo con ese título."}
        
    # extraemos la primera fila que coincida (iloc[0]) y la convertimos a un diccionario
    articulo_dict = filtro.iloc[0].to_dict()
    
    # Limpiamos los valores NaN para q  no se vea feo
    for clave, valor in articulo_dict.items():
        if pd.isna(valor):
            articulo_dict[clave] = "No disponible"
            
    return articulo_dict

#--------------------------------------------------------------------------------------------

def calcular_proporcion_citadas(df):
    """
    Calcula el porcentaje de publicaciones que tienen al menos 1 cita.
    (Proporción de Publicaciones Citadas - PCP)
    """
    col_citas = next((c for c in df.columns if 'cite' in c.lower() or 'cit' in c.lower()), None)
            
    if col_citas is None:
        return {"error": "No se encontró la columna de Citas en el dataset."}
        
    # Forzamos la conversión a números y los valores corruptos/vacíos los volvemos 0
    citas_numericas = pd.to_numeric(df[col_citas], errors='coerce').fillna(0)
    
    total_articulos = len(df)
    if total_articulos == 0:
        return {"error": "El dataset está vacío."}
        
    # Magia de Pandas: (citas_numericas > 0) crea una lista de True/False.
    # Al sumarla, cuenta todos los 'True' (artículos con 1 o más citas).
    articulos_citados = int((citas_numericas > 0).sum())
    
    # Calculamos el porcentaje
    proporcion = (articulos_citados / total_articulos) * 100
    
    return {
        "total_articulos": total_articulos,
        "articulos_citados": articulos_citados,
        "proporcion_pcp": round(proporcion, 2),
        "mensaje_formateado": f"{articulos_citados} de {total_articulos} artículos han sido citados ({round(proporcion, 2)}%)"
    }

#--------------------------------------------------------------------------------------------



#--------------------------------------------------------------------------------------------



#--------------------------------------------------------------------------------------------



#--------------------------------------------------------------------------------------------



#--------------------------------------------------------------------------------------------


#--------------------------------------------------------------------------------------------

def obtener_articulo_sin_citas(df):
 
    col_citas = next((c for c in df.columns if 'cite' in c.lower()), None)
    if col_citas:
        serie = pd.to_numeric(df[col_citas], errors='coerce')
        return int((serie == 0).sum())   # NaN no cuenta como 0
    return 0

def obtener_top_citas_anuales(df, top=10):
    """Calcula el promedio de citas anuales"""
    import datetime
    anio_actual = datetime.datetime.now().year
    
    df = df.copy()
    col_citas = next((c for c in df.columns if 'cite' in c.lower()), None)
    col_anios = next((c for c in df.columns if 'year' in c.lower() or 'año' in c.lower()), None)
    
    # Búsqueda dinámica del título
    col_titulo = next((c for c in df.columns if c.lower() in ['título', 'title', 'ti', 'article title']), 'Title')

    if col_citas and col_anios:
        df[col_citas] = pd.to_numeric(df[col_citas], errors = 'coerce').fillna(0)
        df[col_anios] = pd.to_numeric(df[col_anios], errors = 'coerce')

        df['antiguedad'] = (anio_actual - df[col_anios] + 1).clip(lower=1)
        df['promedio_citas'] = df[col_citas] / df['antiguedad']

        top_df = df.sort_values(by = 'promedio_citas', ascending = False).head(top)
        return top_df[[col_titulo, 'promedio_citas']].to_dict(orient = 'records')
    return []

def calcular_tasa_crecimiento(df):
    col_anio = next((c for c in df.columns if 'year' in c.lower() or 'año' in c.lower()), None)
    if not col_anio:
        return []
    conteo_anual = df[col_anio].value_counts().sort_index()
    tasa = conteo_anual.pct_change()*100

    resultado = []
    for anio, valor in tasa.items():
        resultado.append({
            "año": int(anio),
            "crecimiento": round(valor, 2) if pd.notnull(valor) else 0 
        })
    return resultado

def distribucion_idiomas(df):
    col_idioma = next((c for c in df.columns if 'lang' in c.lower() or 'idioma' in c.lower()), None)
    if not col_idioma:
        return {}
    distribucion = df[col_idioma].value_counts().to_dict()
    return distribucion
#--------------------------------------------------------------------------------------------
  
    
def excel_descargar(df, resumen):
    """
    Exporta un Excel con una hoja por métrica analizada.
    resumen es el diccionario que devuelve procesar_bibliometria()
    """
    output = io.BytesIO()

    with pd.ExcelWriter(output, engine='openpyxl') as writer:

        # Hoja 1 — Métricas generales
        metricas = resumen.get('metricas', {})
        pd.DataFrame([metricas]).to_excel(writer, sheet_name='Metricas_Generales', index=False)

        # Hoja 2 — Top 10 trabajos más citados con APA
        citados = resumen.get('tops', {}).get('citados', [])
        if citados:
            pd.DataFrame(citados).to_excel(writer, sheet_name='Top_Citados', index=False)

        # Hoja 3 — Top 10 autores
        top_autores = resumen.get('analisis_avanzado', {}).get('top_10', [])
        if top_autores:
            pd.DataFrame(top_autores).to_excel(writer, sheet_name='Top_Autores', index=False)

        # Hoja 4 — Universidades
        univs = resumen.get('Afiliaciones', {}).get('Universidades', [])
        if univs:
            pd.DataFrame(univs, columns=['Universidad', 'Frecuencia']).to_excel(
                writer, sheet_name='Universidades', index=False)

        # Hoja 5 — Países
        paises = resumen.get('tops', {}).get('paises', [])
        if paises:
            pd.DataFrame(paises, columns=['País', 'Frecuencia']).to_excel(
                writer, sheet_name='Paises', index=False)

        # Hoja 6 — Revistas
        revistas = resumen.get('tops', {}).get('revistas', [])
        if revistas:
            pd.DataFrame(revistas, columns=['Revista', 'Artículos']).to_excel(
                writer, sheet_name='Top_Revistas', index=False)

        # Hoja 7 — Autor único
        autor_unico = resumen.get('autor_unico', {}).get('articulos', [])
        if autor_unico:
            pd.DataFrame(autor_unico, columns=['Título', 'Autor']).to_excel(
                writer, sheet_name='Autor_Unico', index=False)

        # Hoja 8 — Dataset completo (por si lo necesitan)
        df.to_excel(writer, sheet_name='Dataset_Original', index=False)

    output.seek(0)
    return output
#--------------------------------------------------------------------------------------------

def word_descargar(df, resumen, titulo="Reporte de Análisis Bibliométrico"):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.add_heading(titulo, 0)

    # ── Métricas generales ──────────────────────────────────────
    doc.add_heading('1. Métricas Generales', level=1)
    m = resumen.get('metricas', {})
    aa = resumen.get('analisis_avanzado', {})

    tabla_m = doc.add_table(rows=1, cols=2)
    tabla_m.style = 'Table Grid'
    tabla_m.rows[0].cells[0].text = 'Métrica'
    tabla_m.rows[0].cells[1].text = 'Valor'

    metricas_mostrar = [
        ('Total de artículos',           m.get('total_articulos', '')),
        ('Artículos con título',          m.get('articulos_validos', '')),
        ('Artículos sin citas',           m.get('articulos_sin_citas', '')),
        ('Artículos de autor único',      m.get('total_autor_unico', '')),
        ('Mínimo de autores/artículo',   m.get('min_autores', '')),
        ('Máximo de autores/artículo',   m.get('max_autores', '')),
        ('Promedio de autores/artículo', m.get('promedio_autores', '')),
        ('Promedio de citas anual',       resumen.get('promedio_citas_anual', '')),
        ('Rango de años',                aa.get('rango', {}).get('mensaje_formateado', '')),
    ]
    for etiqueta, valor in metricas_mostrar:
        fila = tabla_m.add_row().cells
        fila[0].text = etiqueta
        fila[1].text = str(valor)

    doc.add_paragraph()

    # ── Top 10 trabajos más citados ─────────────────────────────
    doc.add_heading('2. Top 10 Trabajos Más Citados (formato APA)', level=1)
    citados = resumen.get('tops', {}).get('citados', [])
    for i, art in enumerate(citados, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. ").bold = True
        p.add_run(art.get('apa', art.get('titulo', '')))
        p.add_run(f"  [{art.get('citas', 0)} citas]").bold = True

    doc.add_paragraph()

    # ── Top 10 autores ──────────────────────────────────────────
    doc.add_heading('3. Top 10 Autores', level=1)
    top_autores = aa.get('top_10', [])
    tabla_a = doc.add_table(rows=1, cols=2)
    tabla_a.style = 'Table Grid'
    tabla_a.rows[0].cells[0].text = 'Autor'
    tabla_a.rows[0].cells[1].text = 'Artículos'
    for a in top_autores:
        fila = tabla_a.add_row().cells
        fila[0].text = str(a.get('autor', ''))
        fila[1].text = str(a.get('cantidad', ''))

    doc.add_paragraph()

    # ── Universidades ───────────────────────────────────────────
    doc.add_heading('4. Top 10 Universidades', level=1)
    univs = resumen.get('Afiliaciones', {}).get('Universidades', [])
    if univs:
        tabla_u = doc.add_table(rows=1, cols=2)
        tabla_u.style = 'Table Grid'
        tabla_u.rows[0].cells[0].text = 'Universidad'
        tabla_u.rows[0].cells[1].text = 'Apariciones'
        for u in univs:
            fila = tabla_u.add_row().cells
            fila[0].text = str(u[0])
            fila[1].text = str(u[1])

    doc.add_paragraph()

    # ── Países ──────────────────────────────────────────────────
    doc.add_heading('5. Top 10 Países', level=1)
    paises = resumen.get('tops', {}).get('paises', [])
    if paises:
        tabla_p = doc.add_table(rows=1, cols=2)
        tabla_p.style = 'Table Grid'
        tabla_p.rows[0].cells[0].text = 'País'
        tabla_p.rows[0].cells[1].text = 'Apariciones'
        for p in paises:
            fila = tabla_p.add_row().cells
            fila[0].text = str(p.get('pais', ''))
            fila[1].text = str(p.get('cantidad', ''))

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

#--------------------------------------------------------------------------------------------
def distribucion_documentos(df):
    col_tipo = next(
        (c for c in df.columns if 'document type' in c.lower() or 'type' in c.lower()),
        None
    )                                                   
    if col_tipo:
        return df[col_tipo].value_counts().to_dict()
    return {"error": "Columna no encontrada"}

def calcular_h_index(df):
    col_citas = next((c for c in df.columns if 'cite' in c.lower()), None)  
    if col_citas is None:
        return 0
    citas = df[col_citas].fillna(0).astype(int).sort_values(ascending=False).tolist()
    h = 0
    for i, n_citas in enumerate(citas):
        if n_citas >= i + 1:
            h = i + 1
        else:
            break
    return h

def identificar_tendencias(df):
    try:
        col_citas = next((c for c in df.columns if 'cite' in c.lower()), None)   
        col_anio  = next((c for c in df.columns if 'year' in c.lower()), None)   

        if col_citas is None or col_anio is None:
            return []

        df = df.copy()
        anio_actual = datetime.datetime.now().year            

        col_titulo = next((c for c in df.columns if c.lower() in ['título', 'title', 'ti', 'article title']), 'Title')

        df['Edad'] = (anio_actual - pd.to_numeric(df[col_anio], errors='coerce')).clip(lower=1)
        df['Crecimiento'] = pd.to_numeric(df[col_citas], errors='coerce').fillna(0) / df['Edad']

        cols = [c for c in [col_titulo, 'Authors', 'Crecimiento'] if c in df.columns]
        top = df.sort_values(by='Crecimiento', ascending=False).head(5)
        return top[cols].to_dict(orient='records')
    except Exception as e:
        print(f"Error en identificar_tendencias: {e}")
        return []
    
                              

def generar_grafo_palabras(df):
    col_kw = next(
        (c for c in df.columns if 'keyword' in c.lower()),
        None
    )                                                   #  detecta dinámicamente 
    if col_kw is None:
        return {"nodes": [], "edges": []}

    G = nx.Graph()
    keywords_list = df[col_kw].dropna().str.split(';')

    for tags in keywords_list:
        tags = [t.strip().lower() for t in tags if t.strip()]
        for i in range(len(tags)):
            for j in range(i + 1, len(tags)):
                if G.has_edge(tags[i], tags[j]):
                    G[tags[i]][tags[j]]['weight'] += 1
                else:
                    G.add_edge(tags[i], tags[j], weight=1)

    nodes = [{"id": node, "size": G.degree(node)} for node in G.nodes()]      
    edges = [{"source": u, "target": v, "value": d['weight']} for u, v, d in G.edges(data=True)]
    return {"nodes": nodes, "edges": edges}