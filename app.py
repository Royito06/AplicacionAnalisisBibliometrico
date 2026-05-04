import os
import io
import pandas as pd
from flask import Flask, request, jsonify, render_template, send_file
from docx import Document
from docx.shared import Pt
import src.metrics
import src.cleaner

app = Flask(__name__)
app.secret_key = "bibliometria_total_export_2026"
ultimo_df = None

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload():
    global ultimo_df
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No hay archivo"}), 400
            
        file = request.files['file']
        df = pd.read_excel(file, engine='openpyxl') if file.filename.endswith(('.xlsx', '.xls')) else pd.read_csv(file)

        df.columns = df.columns.str.strip()
        df = src.cleaner.limpiar_dataset(df)
        ultimo_df = df.copy()

        mapeo = {
            'titulo': ['title', 'titulo', 'ti', 'article title'],
            'citas': ['citations', 'citas', 'tc', 'cite', 'times cited', 'cited by'],
            'anio': ['year', 'anio', 'año', 'py', 'publication year'],
            'autor': ['authors', 'author', 'autor', 'au']
        }

        def detectar(claves):
            return next((c for c in df.columns if any(k == c.lower() for k in claves)), None)

        c_titulo = detectar(mapeo['titulo'])
        c_citas = detectar(mapeo['citas'])
        c_anio = detectar(mapeo['anio'])
        c_autor = detectar(mapeo['autor'])

        if not c_titulo or not c_citas:
            return jsonify({"error": "No se detectaron columnas de Título o Citas"}), 400

        citas_serie = pd.to_numeric(df[c_citas], errors='coerce').fillna(0)
        anios = pd.to_numeric(df[c_anio], errors='coerce').dropna() if c_anio else []
        
        return jsonify({
            "rango_anios": {"mensaje_formateado": f"{int(min(anios))}-{int(max(anios))}" if len(anios)>0 else "S/D"},
            "total_citas": {"mensaje_formateado": str(int(citas_serie.sum()))},
            "impacto_citas": {"mensaje_formateado": str(round(citas_serie.mean(), 2))},
            "proporcion_citadas": {"mensaje_formateado": f"{((citas_serie > 0).mean() * 100):.1f}%"},
            "analisis_avanzado": {
                "top_10": [{"Author": k, "Count": int(v)} for k, v in df[c_autor].str.split(';').explode().str.strip().value_counts().head(10).items()] if c_autor else [],
                "top_trabajos": [{"titulo": str(row[c_titulo]), "citas": int(row[c_citas])} for _, row in df.sort_values(by=c_citas, ascending=False).head(10).iterrows()]
            }
        }), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/exportar/excel')
def export_excel():
    if ultimo_df is None: return "Sin datos", 400
    out = io.BytesIO()
    ultimo_df.to_excel(out, index=False)
    out.seek(0)
    return send_file(out, as_attachment=True, download_name="analisis_bibliometrico.xlsx")

@app.route('/exportar/word')
def export_word():
    if ultimo_df is None: return "Sin datos", 400
    
    doc = Document()
    doc.add_heading('Reporte Bibliométrico Integral', 0)
    doc.add_paragraph(f"Total de registros exportados: {len(ultimo_df)}")

    table = doc.add_table(rows=1, cols=len(ultimo_df.columns))
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    for i, col in enumerate(ultimo_df.columns):
        hdr_cells[i].text = str(col)

    for _, row in ultimo_df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val) if pd.notna(val) else ""
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return send_file(out, as_attachment=True, download_name="reporte_completo.docx")

@app.route('/exportar/paises')
def export_paises():
    if ultimo_df is None: return "Sin datos", 400
    paises = src.metrics.obtener_lista_paises(ultimo_df)
    df_paises = pd.DataFrame(paises)
    out = io.BytesIO()
    df_paises.to_excel(out, index=False)
    out.seek(0)
    return send_file(out, as_attachment=True, download_name="distribucion_geografica.xlsx")

if __name__ == '__main__':
    app.run(debug=True, port=5000)